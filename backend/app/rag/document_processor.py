import os
import io
import uuid
import logging
from datetime import datetime
import fitz  # PyMuPDF
import camelot
import pytesseract
from PIL import Image
from docx import Document
import tiktoken
from typing import List, Dict, Any, Optional
from langchain.text_splitter import RecursiveCharacterTextSplitter
from config import Config

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Process documents and create chunks with metadata"""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=Config.CHUNK_SIZE,
            chunk_overlap=Config.CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        self.tokenizer = tiktoken.get_encoding("cl100k_base")
    
    def process_document(self, file_path: str, file_metadata: Dict) -> List[Dict]:
        """Process a document based on its type and return chunks"""
        try:
            mime_type = file_metadata['mime_type']
            file_name = file_metadata['file_name']
            
            logger.info(f"Processing {file_name} ({mime_type})")
            
            if mime_type == 'application/pdf':
                return self._process_pdf(file_path, file_metadata)
            elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                              'application/msword']:
                return self._process_docx(file_path, file_metadata)
            elif mime_type == 'text/plain':
                return self._process_txt(file_path, file_metadata)
            else:
                logger.warning(f"Unsupported file type: {mime_type}")
                return []
                
        except Exception as e:
            logger.error(f"Failed to process document {file_path}: {e}")
            return []
    
    def _process_pdf(self, file_path: str, file_metadata: Dict) -> List[Dict]:
        """Process PDF file and extract text, tables, and images"""
        chunks = []
        
        try:
            doc = fitz.open(file_path)
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Extract tables first
                table_chunks = self._extract_pdf_tables(file_path, page_num + 1, file_metadata)
                chunks.extend(table_chunks)
                
                # Extract images
                image_chunks = self._extract_pdf_images(page, page_num + 1, file_metadata)
                chunks.extend(image_chunks)
                
                # Extract text and chunk it
                text = page.get_text()
                if text.strip():
                    text_chunks = self._chunk_text(text, page_num + 1, file_metadata)
                    chunks.extend(text_chunks)
            
            doc.close()
            logger.info(f"Extracted {len(chunks)} chunks from PDF")
            
        except Exception as e:
            logger.error(f"Failed to process PDF {file_path}: {e}")
        
        return chunks
    
    def _extract_pdf_tables(self, file_path: str, page_num: int, file_metadata: Dict) -> List[Dict]:
        """Extract tables from PDF page using camelot"""
        chunks = []
        
        try:
            # Extract tables from specific page
            tables = camelot.read_pdf(file_path, pages=str(page_num))
            
            for table_idx, table in enumerate(tables):
                if not table.df.empty:
                    # Convert table to markdown format
                    table_markdown = table.df.to_markdown(index=False)
                    
                    chunk = self._create_chunk(
                        text=table_markdown,
                        file_metadata=file_metadata,
                        page_num=page_num,
                        chunk_index=table_idx,
                        is_table=True
                    )
                    chunks.append(chunk)
            
        except Exception as e:
            logger.debug(f"No tables found on page {page_num}: {e}")
        
        return chunks
    
    def _extract_pdf_images(self, page, page_num: int, file_metadata: Dict) -> List[Dict]:
        """Extract images from PDF page and perform OCR"""
        chunks = []
        
        try:
            image_list = page.get_images()
            
            for img_idx, img in enumerate(image_list):
                try:
                    # Extract image
                    xref = img[0]
                    pix = fitz.Pixmap(page.parent, xref)
                    
                    if pix.n - pix.alpha < 4:  # Only process if not CMYK
                        # Convert to PIL Image
                        img_data = pix.tobytes("png")
                        image = Image.open(io.BytesIO(img_data))
                        
                        # Perform OCR
                        ocr_text = pytesseract.image_to_string(image)
                        
                        if ocr_text.strip():
                            chunk = self._create_chunk(
                                text=ocr_text,
                                file_metadata=file_metadata,
                                page_num=page_num,
                                chunk_index=img_idx,
                                is_image=True,
                                ocr_text=ocr_text
                            )
                            chunks.append(chunk)
                    
                    pix = None  # Clean up
                    
                except Exception as e:
                    logger.debug(f"Failed to process image {img_idx} on page {page_num}: {e}")
                    continue
            
        except Exception as e:
            logger.debug(f"No images found on page {page_num}: {e}")
        
        return chunks
    
    def _process_docx(self, file_path: str, file_metadata: Dict) -> List[Dict]:
        """Process DOCX file and extract text and tables"""
        chunks = []
        
        try:
            doc = Document(file_path)
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    # Convert to markdown
                    header = " | ".join(table_data[0])
                    separator = " | ".join(["-"] * len(table_data[0]))
                    rows = [" | ".join(row) for row in table_data[1:]]
                    table_markdown = "\n".join([header, separator] + rows)
                    
                    chunk = self._create_chunk(
                        text=table_markdown,
                        file_metadata=file_metadata,
                        page_num=1,  # DOCX doesn't have clear page numbers
                        chunk_index=table_idx,
                        is_table=True
                    )
                    chunks.append(chunk)
            
            # Extract text from paragraphs
            text_content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            if text_content.strip():
                text_chunks = self._chunk_text(text_content, 1, file_metadata)
                chunks.extend(text_chunks)
            
            logger.info(f"Extracted {len(chunks)} chunks from DOCX")
            
        except Exception as e:
            logger.error(f"Failed to process DOCX {file_path}: {e}")
        
        return chunks
    
    def _process_txt(self, file_path: str, file_metadata: Dict) -> List[Dict]:
        """Process TXT file"""
        chunks = []
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
            
            if content.strip():
                chunks = self._chunk_text(content, 1, file_metadata)
            
            logger.info(f"Extracted {len(chunks)} chunks from TXT")
            
        except Exception as e:
            logger.error(f"Failed to process TXT {file_path}: {e}")
        
        return chunks
    
    def _chunk_text(self, text: str, page_num: int, file_metadata: Dict) -> List[Dict]:
        """Split text into chunks using LangChain splitter"""
        chunks = []
        
        try:
            text_chunks = self.text_splitter.split_text(text)
            
            for idx, chunk_text in enumerate(text_chunks):
                if chunk_text.strip():
                    chunk = self._create_chunk(
                        text=chunk_text,
                        file_metadata=file_metadata,
                        page_num=page_num,
                        chunk_index=idx
                    )
                    chunks.append(chunk)
        
        except Exception as e:
            logger.error(f"Failed to chunk text: {e}")
        
        return chunks
    
    def _create_chunk(self, text: str, file_metadata: Dict, page_num: int, 
                     chunk_index: int, is_table: bool = False, is_image: bool = False, 
                     ocr_text: Optional[str] = None) -> Dict:
        """Create a chunk with metadata"""
        
        # Count tokens
        token_count = len(self.tokenizer.encode(text))
        
        # Determine document type from file extension
        file_name = file_metadata['file_name']
        if file_name.endswith('.pdf'):
            doc_type = 'pdf'
        elif file_name.endswith(('.docx', '.doc')):
            doc_type = 'docx'
        elif file_name.endswith('.txt'):
            doc_type = 'txt'
        else:
            doc_type = 'unknown'
        
        chunk = {
            'chunk_id': str(uuid.uuid4()),
            'text': text,
            'source_doc_id': file_metadata['file_id'],
            'source_doc_name': file_metadata['file_name'],
            'folder_type': file_metadata['parent_folder_type'],
            'owner_email': file_metadata['owner_email'],
            'uploaded_by': 'ingest_service',
            'doc_type': doc_type,
            'created_at': file_metadata['created_at'],
            'ingested_at': str(datetime.now().isoformat()),
            'chunk_index': chunk_index,
            'source_page': page_num,
            'language': 'en',
            'token_count': token_count,
            'is_table': is_table,
            'is_image': is_image,
            'doc_url': f"https://drive.google.com/file/d/{file_metadata['file_id']}/view"
        }
        
        if ocr_text:
            chunk['ocr_text'] = ocr_text
        
        # Will be filled in PHASE 3
        chunk['is_pi'] = None
        chunk['access_roles'] = []
        chunk['uid'] = file_metadata.get('uid')
        
        return chunk

# Test function
def test_document_processing():
    """Test document processing with a sample file"""
    processor = DocumentProcessor()
    
    # Create a test file
    test_file = os.path.join(Config.TEMP_DIR, 'test_document.txt')
    os.makedirs(Config.TEMP_DIR, exist_ok=True)
    
    with open(test_file, 'w') as f:
        f.write("This is a test document.\n\n")
        f.write("It contains multiple paragraphs to test chunking.\n\n")
        f.write("Each paragraph should be processed correctly.")
    
    # Mock file metadata
    file_metadata = {
        'file_id': 'test_123',
        'file_name': 'test_document.txt',
        'mime_type': 'text/plain',
        'created_at': '2024-01-01T00:00:00Z',
        'owner_email': 'test@example.com',
        'parent_folder_type': 'NON_PI',
        'uid': None
    }
    
    chunks = processor.process_document(test_file, file_metadata)
    
    print(f"Generated {len(chunks)} chunks:")
    for i, chunk in enumerate(chunks):
        print(f"\nChunk {i+1}:")
        print(f"- ID: {chunk['chunk_id']}")
        print(f"- Text: {chunk['text'][:100]}...")
        print(f"- Tokens: {chunk['token_count']}")
        print(f"- Type: {chunk['doc_type']}")
    
    # Clean up
    os.remove(test_file)
    return chunks

if __name__ == "__main__":
    from datetime import datetime
    test_document_processing()